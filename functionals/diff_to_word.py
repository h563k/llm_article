import difflib
import os
from docx import Document
from docx.shared import RGBColor
from functionals.system_config import ModelConfig
import re


def split_paragraphs(text):
    # 按空行或换行符分割段落（需适配用户的实际段落分隔逻辑）
    return [p.strip() for p in text.split('\n') if p.strip()]


def apply_diff_with_format(para1, para2, doc_paragraph):
    # 使用正则表达式提取标签和文本
    pattern = re.compile(r'(<[^>]+>)|([^<]+)')
    tokens1 = pattern.findall(para1)
    tokens2 = pattern.findall(para2)

    # 对比纯文本部分（忽略标签）
    text1 = ''.join([t[1] for t in tokens1 if t[1]])
    text2 = ''.join([t[1] for t in tokens2 if t[1]])
    matcher = difflib.SequenceMatcher(None, text1, text2)

    # 生成带格式的差异段落
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        # 处理标签与文本的映射关系
        # （此处需实现标签与文本位置的精确匹配，篇幅限制略去具体实现）
        # 标记差异部分
        if tag == 'replace' or tag == 'delete':
            run = doc_paragraph.add_run(text1[i1:i2])
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.strike = True
        if tag == 'replace' or tag == 'insert':
            run = doc_paragraph.add_run(text2[j1:j2])
            run.font.color.rgb = RGBColor(0, 255, 0)
        if tag == 'equal':
            run = doc_paragraph.add_run(text1[i1:i2])


def word_diff(text1, text2, doc):
    paras1 = split_paragraphs(text1)
    paras2 = split_paragraphs(text2)

    # 确保段落对齐（示例假设段落数一致）
    for p1, p2 in zip(paras1, paras2):
        # 添加原段落
        para_original = doc.add_paragraph()
        apply_diff_with_format(p1, p2, para_original)

        # # 添加润色段落
        # para_edited = doc.add_paragraph()
        # apply_diff_with_format(p2, p1, para_edited)  # 反向对比以显示新增内容


def create_word(text_list, file_name):
    config = ModelConfig()
    file_path = config.file_path
    print(file_path)
    # 创建 Word 文档
    doc = Document()

    # 标题
    doc.add_heading(f'{file_name}文本比较版', 0)

    for text1, text2 in text_list:
        word_diff(text1, text2, doc)
    save_path = os.path.join(file_path, 'data', 'word_diff', file_name)

    doc.save(save_path)
