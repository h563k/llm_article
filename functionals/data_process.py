import os
import re
from functionals.system_config import ModelConfig
from functionals.save_data import DatabaseManager
from functionals.standard_log import log_to_file
from functionals.doc2docx import convert_doc_to_docx
from docx import Document


@log_to_file
def extract_text_with_format(doc_path):
    doc = Document(doc_path)
    formats = []
    for paragraph in doc.paragraphs:
        # 处理段落样式
        style_tag = ""
        if paragraph.style.name.startswith('Heading'):
            style_tag = f"<h{paragraph.style.name[-1]}>"
            close_tag = f"</h{paragraph.style.name[-1]}>"
        # 处理段落内格式
        format_text = ""
        for run in paragraph.runs:
            text = run.text
            if text:
                # 格式包装
                formats_stack = []
                # if run.bold:
                #     formats_stack.append("b")
                if run.italic:
                    formats_stack.append("i")
                if run.underline:
                    formats_stack.append("u")
                if run.font.subscript:
                    formats_stack.append("sub")
                if run.font.superscript:
                    formats_stack.append("sup")
                for tag in formats_stack:
                    text = f"<{tag}>{text}</{tag}>"
            format_text += text
        # format_text = format_text.replace(f'</b><b>', '')
        format_text = format_text.replace(f'</i><i>', '')
        format_text = format_text.replace(f'</u><u>', '')
        format_text = format_text.replace(f'</sub><sub>', '')
        format_text = format_text.replace(f'</sup><sup>', '')
        # 添加段落样式
        if style_tag:
            format_text = f"{style_tag}{format_text}{close_tag}"
        formats.append(format_text)
    # 处理表格
    table_list = []
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            table_content.append("|".join(cells))
        table_list.append("<table>" + "\\n".join(table_content) + "</table>")

    return "\n".join(formats), table_list


def content_process(content, title):
    content_set_list = [title]
    title_list = re.findall(
        '(<h\d+>.*?</h\d+>|\n\d+\x20*.*|\n\d+\.\d+\x20*.*)', content)
    content_set = re.search(
        f'(.*?){re.escape(title_list[0])}', content, re.S).group(1)
    if content_set:
        content_set_list.append(content_set)
        content = content.replace(content_set, '')
    for i in range(len(title_list) - 1):
        content_set = re.search(
            f'(.*){re.escape(title_list[i + 1])}', content, re.S)
        if content_set:
            content_set_list.append(content_set.group(1))
            content = content.replace(content_set.group(1), '')
    content_set = re.search(f'({re.escape(title_list[-1])}.*)', content, re.S)
    content_set_list.append(content_set.group(1))
    return content_set_list


def DataProcess():
    config = ModelConfig()
    DataBase = DatabaseManager()
    WorkSpace = config.WorkSpace
    original_data = WorkSpace['original_data']
    original_data = os.path.join(config.file_path, original_data)
    print(f'原始数据路径为{original_data}')
    log_list = []
    _, _, paths = os.walk(original_data).__next__()
    for path in paths:
        if path.startswith('~$'):
            os.remove(os.path.join(original_data, path))
            print(f'{path}为临时文件, 删除')
            continue
        if path.endswith('.doc'):
            print(f'{path}为doc文件, 开始转换文件')
            doc_path = os.path.join(original_data, path)
            docx_path = f'{doc_path}x'
            convert_doc_to_docx(doc_path, docx_path)
            continue
        if DataBase.get_document(path):
            print(f'{path}已存在, 跳过预处理')
            continue
        file_path = os.path.join(original_data, path)
        print(f'正在处理{path}')
        page_content, table_list = extract_text_with_format(file_path)
        FullContent = page_content
        DataBase.save_document(path, 'FullContent', page_content)
        print(len(page_content))
        check = 0
        for content in page_content.split('\n'):
            if re.findall('(DOI.*)', content, re.I):
                doi = re.findall('(DOI.*)', content, re.I)[0]
                page_content = page_content.replace(doi, '')
                page_content = page_content.strip('\n')
                check += len(doi)
                #  储存 DOI
                DataBase.save_document(path, 'DOI', doi)
                break
        title = re.search('(.*?)摘(\x20*|\u3000*)要',
                          page_content, re.S).group(1)
        # 储存标题
        check += len(title)
        page_content = page_content.replace(title, '')
        abstract = re.search('(摘(\x20*|\u3000*)要.*?)0(\.*)(\x20*|\u3000*)(引|前)(\x20*|\u3000*)言',
                             page_content, re.S).group(1)
        DataBase.save_document(path, 'Abstract', abstract)
        abstract_chinese = re.search('(摘(\x20*|\u3000*)要.*关键(词|字).*?\n)',
                                     abstract, re.S).group(1)
        DataBase.save_document(path, 'Abstract_Chinese', abstract_chinese)
        abstract_english = re.search('(Abstract.*?Key(\x20*|\u3000*)word.*?\n)',
                                     abstract, re.S | re.I)
        if not abstract_english:
            abstract_english = re.search('(Abstract.*?Key(\x20*|\u3000*)word.*?\n)',
                                         FullContent, re.S | re.I).group(1)
        else:
            abstract_english = abstract_english.group(1)
        DataBase.save_document(path, 'Abstract_English', abstract_english)
        # 储存摘要
        check += len(abstract)
        page_content = page_content.replace(abstract, '')
        content = re.search('(0(\.*)(\x20*|\u3000*)(引|前)(\x20*|\u3000*)言.*?)参考文献',
                            page_content, re.S).group(1)
        content_set_list = content_process(content, title)
        # 储存正文
        for content_set in content_set_list:
            if not content_set:
                continue
            check += len(content_set)
            DataBase.save_document(path, 'Content', content_set)
        page_content = page_content.replace(content, '')
        check += len(page_content)
        DataBase.save_document(path, 'Reference', page_content)
        txt = f'全文长为: {check - len(page_content)}, 参考文献长为: {len(page_content)}, 正文分割长度为: {[len(x) for x in content_set_list]}'
        print(txt)
        # 储存表格
        for table in table_list:
            check += len(table)
            DataBase.save_document(path, 'Table', table)
        txt += f'表格长为: {[len(x) for x in table_list]}'
        log_list.append(txt)
    return log_list
